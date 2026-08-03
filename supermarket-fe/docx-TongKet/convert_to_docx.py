import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def md_to_docx(md_path, docx_path):
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip('\n')
        
        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
            continue
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            continue
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
            continue
            
        # Horizontal rule
        if line == '---':
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.add_run('***').bold = True
            continue
            
        # Bullet list
        if line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            text = line[2:]
        else:
            if not line.strip():
                continue
            p = doc.add_paragraph()
            text = line
            
        # Parse bold (**text**) and italic (*text*)
        # Simple regex tokenizer
        tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        for token in tokens:
            if token.startswith('**') and token.endswith('**'):
                run = p.add_run(token[2:-2])
                run.bold = True
            elif token.startswith('*') and token.endswith('*'):
                run = p.add_run(token[1:-1])
                run.italic = True
            else:
                p.add_run(token)

    doc.save(docx_path)

import sys
if len(sys.argv) == 3:
    md_to_docx(sys.argv[1], sys.argv[2])
    print(f"Saved {sys.argv[2]}")
else:
    md_to_docx('CHUNG4_KETLUAN.md', 'CHUNG4_KETLUAN.docx')
    print("Saved CHUNG4_KETLUAN.docx")
